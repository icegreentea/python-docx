import pytest

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.api import Document as OpenDocument
from docx.numbering import (NumberingLevelDefinition,
                            AbstractNumberingDefinition,
                            NumberingInstance,
                            NumberLevelOverride)

from .unitutil.cxml import element, xml


class DescribeNumberingInstance:
    def it_sets_start_override_by_default(self):
        num_elm = element('w:num{w:numId=10}')
        num = NumberingInstance(num_elm, None, None)
        assert 1 == len(num.level_overrides)
        assert 0 == num.level_overrides[0].numbering_level
        assert 1 == num.level_overrides[0].start_override

    def it_knows_its_numbering_instance_id(self):
        num_elm = element('w:num{w:numId=10}')
        num = NumberingInstance(num_elm, None, None)
        assert 10 == num.numbering_id 

    def it_can_add_list_paragraph(self):
        doc = OpenDocument()
        doc._element = element('w:document')
        num_elm = element('w:num{w:numId=10}')
        num = NumberingInstance(num_elm, None, doc._part)
        p = num.add_paragraph(0, "para")

        assert 10 == p.numbering_instance_id
        assert 0 == p.numbering_level
        assert "para" == p.text

    def it_can_add_level_override_and_get_level_override(self):
        doc = OpenDocument()
        doc._element = element('w:document')
        num_elm = element('w:num{w:numId=0}')
        num = NumberingInstance(num_elm, None, doc._part)
        assert 1 == len(num.level_overrides)

        num_override = num.add_level_override(1)
        assert 2 == len(num.level_overrides)
        assert isinstance(num_override, NumberLevelOverride)
        ret_num_override = num.get_level_override_by_ilvl(1)
        assert num_override == ret_num_override

        expected_xml = xml('w:lvlOverride{w:ilvl=1}')
        assert expected_xml == num_override._element.xml


class DescribeNumberLevelOverride:
    def it_knows_and_can_set_numbering_level(self):
        nlo = NumberLevelOverride(element('w:lvlOverride{w:ilvl=2}'), None)
        assert 2 == nlo.numbering_level
        nlo.numbering_level = 0
        assert 0 == nlo.numbering_level
        expected_xml = xml('w:lvlOverride{w:ilvl=0}')
        assert expected_xml == nlo._element.xml

    def it_knows_and_can_set_start_override(self):
        elm = element('w:lvlOverride{w:ilvl=0}/w:startOverride{w:val=1}')
        nlo = NumberLevelOverride(elm, None)
        assert 1 == nlo.start_override
        nlo.start_override = 2
        assert 2 == nlo.start_override
        expected_xml = xml('w:lvlOverride{w:ilvl=0}/w:startOverride{w:val=2}')
        assert expected_xml == nlo._element.xml

    def it_contains_override_level_defintion(self):
        elm = element('w:lvlOverride{w:ilvl=0}')
        nlo = NumberLevelOverride(elm, None)
        assert nlo.override_level_definition is None
        ret = nlo.create_override_level_definition()
        assert isinstance(ret, NumberingLevelDefinition)
        expected_xml = xml('w:lvlOverride{w:ilvl=0}/w:lvl')
        assert expected_xml == nlo._element.xml
        nlo.remove_override_level_definition()
        assert nlo.override_level_definition is None
        expected_xml = xml('w:lvlOverride{w:ilvl=0}')
        assert expected_xml == nlo._element.xml


class DescribeAbstractNumberingDefinition:
    def it_is_a_sequence_of_levels(self, simple_abstract_num_fixture):
        abnum, exp_num_levels, exp_name = simple_abstract_num_fixture
        assert exp_num_levels == len(abnum)
        assert exp_name == abnum.name

        for i, lvl in enumerate(abnum):
            assert isinstance(lvl, NumberingLevelDefinition)
            assert i == lvl.numbering_level

        for i in range(len(abnum)):
            assert isinstance(abnum[i], NumberingLevelDefinition)
            assert i == abnum[i].numbering_level

    def it_can_set_number_formats(self, set_number_format_fixture):
        abnum, set_value, expected_xml = set_number_format_fixture
        abnum.set_level_number_format(set_value)
        assert expected_xml == abnum._element.xml

    def it_can_set_level_text(self, set_level_text_fixture):
        abnum, set_value, expected_xml = set_level_text_fixture
        abnum.set_level_text(set_value)
        assert expected_xml == abnum._element.xml

    @pytest.fixture
    def simple_abstract_num_fixture(self):
        return AbstractNumberingDefinition(
            element(
                   'w:abstractNum{w:abstractNumId=1}/('
                   'w:name{w:val=testabnum},'
                   'w:lvl{w:ilvl=0},'
                   'w:lvl{w:ilvl=1},'
                   'w:lvl{w:ilvl=2},'
                   'w:lvl{w:ilvl=3})'
                   )
        ), 4, "testabnum"

    @pytest.fixture(params=[
        ("bullet", 'w:lvl{w:ilvl=0}/w:numFmt{w:val=bullet},w:lvl{w:ilvl=1}/w:numFmt{w:val=bullet}'),    # noqa
        (["bullet", "decimal"], 'w:lvl{w:ilvl=0}/w:numFmt{w:val=bullet},w:lvl{w:ilvl=1}/w:numFmt{w:val=decimal}'),  # noqa
    ])
    def set_number_format_fixture(self, request):
        set_value, tail_expected_cxml = request.param
        abnum = AbstractNumberingDefinition(element(
            'w:abstractNum{w:abstractNumId=1}/(w:lvl{w:ilvl=0},w:lvl{w:ilvl=1})'
        ))

        expected_xml = xml('w:abstractNum{w:abstractNumId=1}/(%s)' % tail_expected_cxml)
        return abnum, set_value, expected_xml

    @pytest.fixture(params=[
        ("1.", 'w:lvl{w:ilvl=0}/w:lvlText{w:val=1.},w:lvl{w:ilvl=1}/w:lvlText{w:val=1.}'),    # noqa
        (["1.", "2."], 'w:lvl{w:ilvl=0}/w:lvlText{w:val=1.},w:lvl{w:ilvl=1}/w:lvlText{w:val=2.}'),  # noqa
    ])
    def set_level_text_fixture(self, request):
        set_value, tail_expected_cxml = request.param
        abnum = AbstractNumberingDefinition(element(
            'w:abstractNum{w:abstractNumId=1}/(w:lvl{w:ilvl=0},w:lvl{w:ilvl=1})'
        ))

        expected_xml = xml('w:abstractNum{w:abstractNumId=1}/(%s)' % tail_expected_cxml)
        return abnum, set_value, expected_xml


class DescribeNumberingLevelDefinition:
    def it_has_basic_getter_and_setter(
                                       self,
                                       minimal_numbering_level_definition_fixture
                                      ):
        num_def, exp_ilvl = minimal_numbering_level_definition_fixture
        assert num_def.numbering_level == exp_ilvl

        assert num_def.start is None
        num_def.start = 10
        assert num_def.start == 10

        assert num_def.number_format is None
        num_def.number_format = "bullet"
        assert num_def.number_format == "bullet"

        assert num_def.restart_numbering_level is None
        num_def.restart_numbering_level = 3
        assert num_def.restart_numbering_level == 3

        assert num_def.numbering_level_text is None
        num_def.numbering_level_text = "%1"
        assert num_def.numbering_level_text == "%1"

        assert num_def.justification is None
        num_def.justification = WD_ALIGN_PARAGRAPH.RIGHT
        assert num_def.justification == WD_ALIGN_PARAGRAPH.RIGHT

        assert num_def.paragraph_properties is None

        expected_xml = xml('w:lvl{w:ilvl=%s}/'
                           '(w:start{w:val=10},w:numFmt{w:val=bullet},'
                           'w:lvlRestart{w:val=3},'
                           'w:lvlText{w:val=%%1},w:lvlJc{w:val=right})' % exp_ilvl)

        assert num_def._element.xml == expected_xml

    def it_provides_paragraph_properties(
                                       self,
                                       minimal_numbering_level_definition_fixture
                                      ):
        num_def, exp_ilvl = minimal_numbering_level_definition_fixture

        assert num_def.paragraph_properties is None
        paragraph_props = num_def.create_new_paragraph_properties()
        paragraph_props.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        assert num_def.paragraph_properties.alignment == WD_ALIGN_PARAGRAPH.RIGHT

        expected_xml = xml("w:lvl{w:ilvl=0}/w:pPr/w:jc{w:val=right}")
        assert num_def._element.xml == expected_xml

        num_def.paragraph_properties.alignment = WD_ALIGN_PARAGRAPH.LEFT
        expected_xml = xml("w:lvl{w:ilvl=0}/w:pPr/w:jc{w:val=left}")
        assert num_def._element.xml == expected_xml

    @pytest.fixture
    def minimal_numbering_level_definition_fixture(self, request):
        return NumberingLevelDefinition(element('w:lvl{w:ilvl=0}')), 0
