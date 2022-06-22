import docx

doc = docx.Document()
ab_num = doc.numbering.create_bullet_abstract_numbering("bullet-list")
num_instance = doc.numbering.create_numbering_instance(ab_num)

p1 = doc.add_paragraph("1.1 Bullet List")
p1.set_numbering(num_instance, 0)

p2 = doc.add_paragraph("2.1 Bullet List")
p2.set_numbering(num_instance, 0)

p3 = doc.add_paragraph("2.2 Bullet List")
p3.set_numbering(num_instance, 1)

p4 = doc.add_paragraph("3.1 Bullet List")
p4.set_numbering(num_instance, 0)

for i in range(0, 9):
    p = doc.add_paragraph(f"{i+1} Bullet List")
    p.set_numbering(num_instance, i)

doc.save("output/bullet_list.docx")