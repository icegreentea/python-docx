import pytest
import os

from docx.api import Document as OpenDocument

CURDIR = os.path.abspath(os.path.dirname(__file__))
OUTPUT_DIR = os.path.join(CURDIR, "output")


def save_document(doc, filename):
    if not os.path.isdir(OUTPUT_DIR):
        os.mkdir(OUTPUT_DIR)
    doc.save(os.path.join(OUTPUT_DIR, filename))


class ManuallyCheckHyperlinks:
    @pytest.mark.manual
    def manually_check_add_external_hyperlink(self):
        """
        Generate file "hyperlinks-1.docx" which contains:
        * Link to `https://www.google.com`
        * Link is styled (blue and underline)
        * Link text is "Link to google"
        * Link is part of overall paragraph with non-hyperlink run elements
        """
        doc = OpenDocument()
        p = doc.add_paragraph()
        p.add_run("Some opening text. ")
        p.add_hyperlink("Link text. ", hyperlink_url="https://www.google.com", document=doc)
        p.add_run("Some more normal text. ")
        save_document(doc, "hyperlinks-1.docx")
