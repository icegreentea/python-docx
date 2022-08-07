from _pytest.compat import num_mock_patch_args
import pytest
import os

import docx
from docx.numbering import AbstractNumberingDefinition, NumberingLevelDefinition

CURDIR = os.path.abspath(os.path.dirname(__file__))
OUTPUT_DIR = os.path.join(CURDIR, "output")


def save_document(doc, filename):
    if not os.path.isdir(OUTPUT_DIR):
        os.mkdir(OUTPUT_DIR)
    doc.save(os.path.join(OUTPUT_DIR, filename))


class ManuallyCheckNumbering:

    @pytest.mark.manual
    def it_can_create_mixed_cases(self):
        """

        """
        doc = docx.Document()
        bullet_abnum = doc.create_new_bullet_definition()
        bullet_num = doc.create_new_numbering_instance(bullet_abnum)
        
        doc.add_paragraph("The following section should contain nested bullet list.")
        for i in range(0,9):
            doc.add_paragraph("Bullet Sweep {}".format(i+1), numbering_instance=bullet_num, indent_level=i)
        doc.add_paragraph("Bullet Reset", numbering_instance=bullet_num, indent_level=0)

        simple_decimal_abnum = doc.create_new_simple_decimal_definition()
        simple_decimal_num = doc.create_new_numbering_instance(simple_decimal_abnum)
        doc.add_paragraph("The following section should contain a simple decimal list.")
        for i in range(0,9):
            doc.add_paragraph("Simple Decimal Sweep {}".format(i+1), 
                numbering_instance=simple_decimal_num, indent_level=i)
        for i in range(0,4):
            doc.add_paragraph("Simple Decimal Count Up", 
                numbering_instance=simple_decimal_num, indent_level=0)

        fully_defined_decimal_abum = doc.create_new_abstract_numbering_definition()
        fully_defined_decimal_abum.set_level_number_format("decimal")
        fully_defined_decimal_abum.set_level_text(AbstractNumberingDefinition.fully_defined_decimal_definition()[1])
        fully_defined_decimal_num = doc.create_new_numbering_instance(fully_defined_decimal_abum)
        doc.add_paragraph("The following section should contain a fully defined decimal list.")
        for i in range(0,9):
            doc.add_paragraph("Fully Defined Decimal Sweep {}".format(i+1), 
                numbering_instance=fully_defined_decimal_num, indent_level=i)
        for i in range(0,4):
            doc.add_paragraph("Fully Defined Decimal Count Up", 
                numbering_instance=fully_defined_decimal_num, indent_level=0)
        for i in range(0,4):
            doc.add_paragraph("Fully Defined Decimal Count Up", 
                numbering_instance=fully_defined_decimal_num, indent_level=1)

        doc.add_paragraph("Now we reset the fully defined decimal count")
        fully_defined_decimal_num2 = doc.create_new_numbering_instance(fully_defined_decimal_abum)
        for i in range(0,9):
            doc.add_paragraph("Fully Defined Decimal Sweep (2) {}".format(i+1), 
                numbering_instance=fully_defined_decimal_num2, indent_level=i)

        save_document(doc, "numbering-spread.docx")

    @pytest.mark.manual
    def it_can_do_numbered_headers(self):
        doc = docx.Document()
        doc.add_heading("HEADING 1")
        
        abnum = doc.create_new_abstract_numbering_definition()
        full_def = AbstractNumberingDefinition.fully_defined_decimal_definition()
        abnum.set_level_number_format(full_def[0])
        abnum.set_level_text(full_def[1])
        for i, lvl in enumerate(abnum):
            lvl : "NumberingLevelDefinition"
            pStyle = lvl._element.get_or_add_pStyle()
            pStyle.val = "Heading {}".format(i+1)
        numist = doc.create_new_numbering_instance(abnum)
        numist.add_paragraph(0, "Heading 1", style="Heading 1")
        numist.add_paragraph(1, "Heading 2", style="Heading 2")


        save_document(doc, "numbered-header-experiment.docx")

    @pytest.mark.manual
    def it_can_create_bullets(self):
        doc = docx.Document()
        abnum = doc.create_new_bullet_definition()
        numist = doc.create_new_numbering_instance(abnum)
        numist.add_paragraph(0, "b1")
        numist.add_paragraph(0, "b2")
        numist.add_paragraph(1, "b3")
        save_document(doc, "bullet-list.docx")

    @pytest.mark.manual
    def it_can_create_alternating_bullets(self):
        doc = docx.Document()
        abnum = doc.create_new_bullet_definition()
        numFmt, lvlTexts = AbstractNumberingDefinition.alternate_bullet_definition()
        abnum.set_level_text(lvlTexts)
        numist = doc.create_new_numbering_instance(abnum)

        for i in range(9):
            numist.add_paragraph(i, "p1")
        save_document(doc, "alternating-bullet-list.docx")

    @pytest.mark.manual
    def it_can_create_fully_defined_decimals(self):
        doc = docx.Document()
        abnum = doc.create_new_bullet_definition()
        numFmt, lvlTexts = \
            AbstractNumberingDefinition.fully_defined_decimal_definition()
        abnum.set_level_number_format(numFmt)
        abnum.set_level_text(lvlTexts)
        numist = doc.create_new_numbering_instance(abnum)

        for i in range(9):
            numist.add_paragraph(i, "p1")
        save_document(doc, "fully-defined-decimal.docx")

    @pytest.mark.manual
    def it_can_create_decimal(self):
        doc = docx.Document()
        abnum = doc.create_new_simple_decimal_definition()
        numist = doc.create_new_numbering_instance(abnum)
        numist.add_paragraph(0, "b1")
        numist.add_paragraph(0, "b2")
        numist.add_paragraph(1, "b3")
        save_document(doc, "numbering-decimal-list.docx")

    @pytest.mark.manual
    def it_can_restart_decimal_lists(self):
        doc = docx.Document()
        abnum = doc.create_new_simple_decimal_definition()
        num_ist_1 = doc.create_new_numbering_instance(abnum)
        num_ist_1.add_paragraph(0, "list-1-para-1")
        num_ist_1.add_paragraph(0, "list-1-para-2")
        num_ist_1.add_paragraph(0, "list-1-para-3")
        doc.add_paragraph("para-break")
        num_ist_2 = doc.create_new_numbering_instance(abnum)
        num_ist_2.add_paragraph(0, "list-2-para-1")
        num_ist_2.add_paragraph(0, "list-2-para-2")
        num_ist_2.add_paragraph(0, "list-2-para-3")
        save_document(doc, "numbering-decimal-list-restart.docx")
